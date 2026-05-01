# -*- coding: utf-8 -*-
"""Generate 测试总结报告.docx (structure aligned with 仿真报告模版)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    return p


def main():
    out = Path(
        r"c:\Users\lenovo\xwechat_files\wxid_rxz24hmaxkka22_cf44\msg\file\2026-05\RISC-V多周期CPU测试总结报告.docx"
    )
    fallback = Path(__file__).resolve().parent.parent / "RISCV_CPU_test_report.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("多周期 RISC-V CPU 功能验证\n测试总结报告")
    r.bold = True
    r.font.size = Pt(22)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("（指令子集：add / sub / and / or / xor / sll / srl / sra / ori / addi / lw / sw / beq / bne / jal / jalr）")
    r2.font.size = Pt(11)

    doc.add_paragraph()

    add_heading(doc, "1. 仿真平台介绍", level=1)
    add_heading(doc, "1.1 仿真工具", level=2)
    add_para(
        doc,
        "RTL 级仿真：Synopsys VCS + Verdi（工程目录 sim/Makefile.vcs，测试平台 tb/sim_vcs.v）。"
        "判定方式：程序进入自旋（PC 稳定）后，读取指定通用寄存器与期望值比较。",
    )
    add_para(
        doc,
        "指令级参考仿真：RARS 1.6（rars1_6.jar）。"
        "访存相关用例采用 mc CompactDataAtZero，其余用例采用 mc CompactTextAtZero，"
        "与数据段布局及教材常用设置一致。",
    )

    add_heading(doc, "1.2 测试环境与通过准则", level=2)
    add_para(
        doc,
        "测试向量：hex 目录下 t01～t38 共 38 组，每组包含同名 .s（汇编注释）与 .hex（机器码）。"
        "刻意未修改 code.c、code.hex。",
    )
    add_para(
        doc,
        "通过准则：① RARS 可汇编（nc a）；② 限定步数内运行结束于自旋；③ 签名寄存器值与 sim_vcs.v 中 is_pass 一致。"
        "VCS 侧与上述寄存器期望值保持一致。",
    )

    add_heading(doc, "2. 测试内容与方法", level=1)
    add_heading(doc, "2.1 指令集与覆盖意图", level=2)
    add_para(
        doc,
        "在 16 条指令范围内，针对译码符号扩展、ALU 边界、移位（含移位量 0/31、sra 正负补位）、"
        "访存（负偏移、写后立刻读、同址覆盖、S/L 型 -2048 边界）、分支（成立/不成立、前向/后向）、"
        "jal/jalr（含 jalr 非零立即数）、x0 恒零（含 lw x0 与 addi 使用 x0 作源）等场景设计用例。",
    )

    add_heading(doc, "2.2 各测试文件说明（逐条）", level=2)
    add_para(
        doc,
        "下表按 t01～t38 列出对应 .s/.hex 文件名、本测例具体验证的行为，以及 tb/sim_vcs.v 中用于判定的签名寄存器与 32 位期望值。",
    )

    # (id, file_stem, detail_zh, pass_expr_zh)
    case_rows = [
        (
            "t01",
            "t01_ralu",
            "R 型 ALU 全覆盖：add/sub/and/or/xor；移位 sll/srl/sra（含对正数 1<<4 再右移）；"
            "负数 addi（-1、-2）相加得 -3，再 addi +3 归零。检验 ALU 运算与寄存器写回。",
            "x15 == 0x00000000",
        ),
        (
            "t02",
            "t02_itype",
            "I 型立即数符号扩展：addi x0,-1 得全 1；+1 归零；ori 与 -256（12 位符号扩展）得 0xFFFFFF00，再 +1 得 0xFFFFFF01。",
            "x4 == 0xFFFFFF01",
        ),
        (
            "t03",
            "t03_mem",
            "lw/sw 基本路径与负偏移：基址 0x20 处 sw/lw 0x123；再在 base-4 写入 0x456 并立即 lw 读回，检验地址计算与写后读。",
            "x5 == 0x00000456",
        ),
        (
            "t04",
            "t04_beq_taken",
            "beq 条件成立时跳转：x1=x2=5 时跳 target，跳过中间 addi x3，最终 x4=1。",
            "x4 == 0x00000001",
        ),
        (
            "t05",
            "t05_beq_not_taken",
            "beq 条件不成立顺序执行：x1≠x2 时执行 ori x3=0x99 与后续 addi x5=1。",
            "x5 == 0x00000001",
        ),
        (
            "t06",
            "t06_bne",
            "bne 成立跳转：跳过错误路径的 addi x6=0x11，在标号处 x6=0x22、x9=0x33、x10=1。",
            "x10 == 0x00000001",
        ),
        (
            "t07",
            "t07_jal_link",
            "jal 保存返回地址：jal x10,target 后 x10=PC+4，跳过 addi x11，目标处 x12=1。",
            "x12 == 0x00000001",
        ),
        (
            "t08",
            "t08_jalr",
            "jalr 按 rs 跳转：目标地址置于 x1（如 0x10），jalr 后应跳过 addi x30（误路径）。"
            "（RARS 下可观察 x5 链路与 PC；VCS 以 x30 未被写为 123 判定跳过成功。）",
            "x30 == 0x00000000",
        ),
        (
            "t09",
            "t09_x0",
            "x0 恒零：尝试 addi x0 与 lw x0 写回后，addi x1,x0,0x66 必须仍得 0x66。",
            "x1 == 0x00000066",
        ),
        (
            "t10",
            "t10_loop_bne",
            "bne 后向循环：x1 从 0 递增至与 x2=3 相等后退出，检验负偏移分支与循环计数。",
            "x1 == 0x00000003",
        ),
        (
            "t11",
            "t11_jal_back",
            "jal x0 前向跳过 NOP，多段顺序执行后 x1=1、x2=2、x3=3，检验无条件跳转与指令流。",
            "x3 == 0x00000003",
        ),
        (
            "t12",
            "t12_sw_neg",
            "sw 负偏移写入 base-28；lw base-4 读到未初始化字（仿真 DM 为 0）。检验地址加法与符号扩展。",
            "x3 == 0x00000000",
        ),
        (
            "t13",
            "t13_shift_boundary",
            "移位边界：移位量 0 与 31；srl 全 1 逻辑右移得 1；sra 负数算术右移保持符号；正数 0x7FF sra 31 得 0。",
            "x13 == 0x00000000",
        ),
        (
            "t14",
            "t14_add_negative",
            "负数与混合符号 add/sub：(-1)+(-2)、(-1)+5、(-1)-5、(-2)-(-1)、16-(-1)、(-1)+1 等，终点 x13=0。",
            "x13 == 0x00000000",
        ),
        (
            "t15",
            "t15_bne_false",
            "bne 不跳转：两对相等寄存器时顺序执行，最终 ori x8=0x88。",
            "x8 == 0x00000088",
        ),
        (
            "t16",
            "t16_lw_negative_extend",
            "控制流组合：bne 后向循环；jal 跳过中间指令；jalr x9,x10,0 将返回地址写入 x9（期望 0x28）。"
            "（文件名偏历史；本条侧重分支与 jal/jalr 配合。）",
            "x9 == 0x00000028",
        ),
        (
            "t17",
            "t17_addi_bounds",
            "addi 立即数边界：-2048 与 +2047 相加得 -1（32 位 0xFFFFFFFF）。",
            "x20 == 0xFFFFFFFF",
        ),
        (
            "t18",
            "t18_ori_and",
            "ori 正立即数 2047 与 ori -1 后按位 and，得 0x7FF。",
            "x20 == 0x000007FF",
        ),
        (
            "t19",
            "t19_xor_sign",
            "-1 与 +1 异或得 0xFFFFFFFE。",
            "x20 == 0xFFFFFFFE",
        ),
        (
            "t20",
            "t20_sub_self",
            "sub 同源：x1-x1=0。",
            "x20 == 0x00000000",
        ),
        (
            "t21",
            "t21_srl_rs0",
            "移位量 rs2=x0（低 5 位为 0）：srl 不改变数据，0x7B 保持。",
            "x20 == 0x0000007B",
        ),
        (
            "t22",
            "t22_sll31",
            "逻辑左移 31 位：1<<31 = 0x80000000。",
            "x20 == 0x80000000",
        ),
        (
            "t23",
            "t23_beq_always",
            "beq x0,x0 恒成立跳过中间 addi x20=1，再写入 0x42。",
            "x20 == 0x00000042",
        ),
        (
            "t24",
            "t24_jal_jalr_ret",
            "子程序调用：jal x5,callee 后 callee 内 x20=0x33，jalr x0,x5,0 返回后执行 addi x21=0xAA。",
            "x20==0x33 且 x21==0xAA",
        ),
        (
            "t25",
            "t25_sw_overwrite",
            "同一字地址连续 sw 0x111 再 0x222，lw 应读到最后一次写入。",
            "x20 == 0x00000222",
        ),
        (
            "t26",
            "t26_alu_chain",
            "连续 RAW：x2=x1+x1，x3=x2+x2，…，x20=x4+x4=16，对流水线/多周期数据前推有压力。",
            "x20 == 0x00000010",
        ),
        (
            "t27",
            "t27_sra_neg1",
            "-1 算术右移 1 位仍为全 1。",
            "x20 == 0xFFFFFFFF",
        ),
        (
            "t28",
            "t28_bne_taken",
            "bne 成立跳过错误 addi x20=0xFF，标号处写入 0x56。",
            "x20 == 0x00000056",
        ),
        (
            "t29",
            "t29_or_combine",
            "ori 与已有模式组合：0xF0 | 0x00F = 0xFF。",
            "x20 == 0x000000FF",
        ),
        (
            "t30",
            "t30_sw_lw_min_offset",
            "S/L 型 12 位有符号立即数下界 -2048：用两条 addi 拼出基址 2048，sw/lw -2048 访问地址 0，读写 0x3EF。",
            "x20 == 0x000003EF",
        ),
        (
            "t31",
            "t31_jalr_imm",
            "jalr 非零立即数：目标 = rs1 + imm；通过 jal 到 L2，再 jalr x0, x10, -4 回到 L1 写入 x20=0xCC。",
            "x20 == 0x000000CC",
        ),
        (
            "t32",
            "t32_add_maxint_wrap",
            "构造 0x7FFFFFFF 后 +1 无异常回绕为 0x80000000（用移位与 xor 组合）。",
            "x20 == 0x80000000",
        ),
        (
            "t33",
            "t33_sra_pos_msb",
            "正数 1<<30 再 sra 31，算术右移后应为 0。",
            "x20 == 0x00000000",
        ),
        (
            "t34",
            "t34_lw_x0_rs",
            "lw 目的寄存器为 x0（不写回），随后 addi x21,x0,5 必须读到 rs1=0，得 x21=5。",
            "x21 == 0x00000005",
        ),
        (
            "t35",
            "t35_mem_chain_offsets",
            "多字 sw/lw 与负偏移搬运：0(x1)、4(x1) 读写后，将一字经 -8(x1) 中转，最终 x20=0x101。",
            "x20 == 0x00000101",
        ),
        (
            "t36",
            "t36_bne_loop10",
            "bne 后向循环 10 次：x20 从 0 加到 10。",
            "x20 == 0x0000000A",
        ),
        (
            "t37",
            "t37_beq_fallthrough2",
            "连续两条 beq 条件均假，必须顺序执行到 addi x20=0x3D，不能误入 join 前缺失该指令。",
            "x20 == 0x0000003D",
        ),
        (
            "t38",
            "t38_sub_from_zero",
            "0 - 0x80000000 的 32 位补码回绕，结果为 0x80000000。",
            "x20 == 0x80000000",
        ),
    ]

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "编号"
    hdr[1].text = "文件（.s / .hex）"
    hdr[2].text = "验证内容（详细）"
    hdr[3].text = "VCS 通过条件"
    for cid, stem, detail, pcond in case_rows:
        row = table.add_row().cells
        row[0].text = cid
        row[1].text = stem
        row[2].text = detail
        row[3].text = pcond

    add_heading(doc, "2.3 自动化与复现", level=2)
    add_para(
        doc,
        "VCS：在 sim 目录执行 make -f Makefile.vcs run TEST=<用例名>，或 make -f Makefile.vcs all 顺序跑完全部 TESTS。",
    )
    add_para(
        doc,
        "RARS：java -jar rars1_6.jar nc mc <CompactTextAtZero|CompactDataAtZero> <步数上限> hex\\<用例>.s x<寄存器> …",
    )

    add_heading(doc, "3. 测试结果与结论", level=1)
    add_heading(doc, "3.1 RARS 指令级仿真结果", level=2)
    add_para(
        doc,
        "在本地对 hex 目录下全部 38 个 .s 用例执行汇编与仿真比对，结果：38/38 通过，"
        "与 tb/sim_vcs.v 中约定的签名寄存器一致。",
        bold=False,
    )

    add_heading(doc, "3.2 VCS / Verdi 结果说明", level=2)
    add_para(
        doc,
        "VCS 需在已正确配置 VERDI_HOME 及授权的环境中编译运行；"
        "若个别用例失败，请对照 simv_<用例>.log 与波形检查 PC 更新、访存地址线（字对齐索引）及寄存器写回使能。",
    )

    add_heading(doc, "3.3 总结", level=2)
    add_para(
        doc,
        "本套测试在 16 条指令约束下，对符号扩展、ALU 与移位边界、访存时序与偏移、控制流与 x0 行为做了分层强化；"
        "RARS 侧已全部验证通过，可作为 RTL 回归基线。后续若扩展指令或修改存储映射，应同步更新 hex 与 is_pass 期望值。",
    )

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = foot.add_run("报告生成说明：结构对齐《仿真报告模版》章节层次；具体内容依据 multicyc 工程当前测试集。")
    fr.font.size = Pt(9)

    try:
        doc.save(out)
        print("saved:", out)
    except PermissionError:
        doc.save(fallback)
        print("saved (fallback, 请关闭 Word 后重试主路径):", fallback)


if __name__ == "__main__":
    main()
